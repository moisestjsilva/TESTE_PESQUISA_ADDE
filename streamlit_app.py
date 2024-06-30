import streamlit as st
from docx import Document
from PIL import Image
import tkinter as tk
from tkinter import filedialog

def select_directory():
    root = tk.Tk()
    root.withdraw()  # Oculta a janela principal do tkinter
    directory = filedialog.askdirectory()
    root.destroy()
    return directory

def main():
    st.title('Aplicativo para Juntar Imagens em Documento DOC')

    # Interface para upload de imagens
    uploaded_files = st.file_uploader("Selecione as imagens que deseja juntar", type=['jpg', 'png'], accept_multiple_files=True)

    # Interface para nomear o arquivo DOC
    doc_name = st.text_input("Nome do arquivo DOC", "meu_documento")

    if st.button("Selecionar Diretório para Salvar"):
        directory = select_directory()
        if directory:
            st.success(f'Diretório selecionado: {directory}')
        else:
            st.error('Nenhum diretório selecionado.')

    if st.button("Criar Documento DOC"):
        if uploaded_files:
            directory = select_directory()
            if not directory:
                st.error('Por favor, selecione um diretório para salvar o documento.')
                return

            doc = Document()

            # Definir margens mínimas (em polegadas)
            margin_top = 0.5  # 0.5 polegadas
            margin_bottom = 0.5  # 0.5 polegadas
            margin_left = 0.5  # 0.5 polegadas
            margin_right = 0.5  # 0.5 polegadas

            # Converter polegadas para pontos (1 polegada = 72 pontos)
            margin_top_in_points = int(margin_top * 72)
            margin_bottom_in_points = int(margin_bottom * 72)
            margin_left_in_points = int(margin_left * 72)
            margin_right_in_points = int(margin_right * 72)

            # Configurar as margens do documento
            sections = doc.sections
            for section in sections:
                section.top_margin = margin_top_in_points
                section.bottom_margin = margin_bottom_in_points
                section.left_margin = margin_left_in_points
                section.right_margin = margin_right_in_points

            # Calcular largura e altura máximas disponíveis no documento
            max_width = doc.sections[0].page_width - margin_left_in_points - margin_right_in_points
            max_height = doc.sections[0].page_height - margin_top_in_points - margin_bottom_in_points

            # Inicializa a largura e altura total das imagens na página
            total_width_on_page = 0
            total_height_on_page = 0

            # Lista para armazenar imagens na linha atual
            current_line_images = []

            for idx, file in enumerate(uploaded_files):
                # Abre a imagem e obtém suas dimensões
                img = Image.open(file)
                width, height = img.size

                # Calcula a proporção da imagem
                aspect_ratio = width / height

                # Calcula a nova largura e altura da imagem para ajustar ao documento
                new_width = max_width
                new_height = int(new_width / aspect_ratio)

                # Verifica se a imagem cabe na linha atual da página
                if total_height_on_page + new_height <= max_height:
                    # Adiciona a imagem à linha atual
                    current_line_images.append((file, new_width, new_height))
                    total_width_on_page = max(total_width_on_page, new_width)
                    total_height_on_page += new_height
                else:
                    # Se não couber, adiciona uma nova linha de imagens
                    add_images_to_document(doc, current_line_images)
                    doc.add_page_break()

                    # Reinicializa para uma nova linha de imagens
                    current_line_images = [(file, new_width, new_height)]
                    total_width_on_page = new_width
                    total_height_on_page = new_height

                # Se for a última imagem, adiciona à linha atual
                if idx == len(uploaded_files) - 1:
                    add_images_to_document(doc, current_line_images)

            # Salva o documento
            doc_path = f"{directory}/{doc_name}.docx"
            doc.save(doc_path)
            st.success(f"Documento {doc_path} criado com sucesso!")

def add_images_to_document(doc, images):
    # Adiciona todas as imagens na lista à linha atual do documento
    for file, width, height in images:
        doc.add_picture(file, width=width, height=height)
        # Adiciona espaço entre imagens (opcional)
        doc.paragraphs[-1].runs[-1].add_break()

if __name__ == '__main__':
    main()
